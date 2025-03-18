#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Módulo controlador para la generación de exámenes en PDF y WORD
"""

import os
import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.platypus.flowables import HRFlowable
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from model.pregunta_dao import PreguntaDAO

class ExamenGenerator:
    """
    Controlador para generar exámenes en formato PDF y Word
    """
    
    def __init__(self):
        """
        Constructor de la clase ExamenGenerator
        """
        self.pregunta_dao = PreguntaDAO()
        self.directorio_examenes = "Examenes"
        
        # Crear directorio si no existe
        if not os.path.exists(self.directorio_examenes):
            os.makedirs(self.directorio_examenes)
    
    def generar_examen_pdf(self, numero_tema):
        """
        Genera una versión de examen en formato PDF
        
        Args:
            numero_tema (int): Número del tema a generar (1, 2, 3, ...)
            
        Returns:
            str: Ruta del archivo PDF generado
        """
        # Generar letra del tema (A, B, C, ...)
        letra_tema = chr(64 + numero_tema)  # 65 es el código ASCII de 'A'
        nombre_archivo = os.path.join(self.directorio_examenes, f"Examen_Tema_{letra_tema}.pdf")
        
        # Generar examen aleatorio
        preguntas_examen = self.pregunta_dao.generar_examen_aleatorio()
        
        # Crear PDF
        if self.generar_pdf(preguntas_examen, nombre_archivo, f"Tema {letra_tema}"):
            return nombre_archivo
        return None
        
    def generar_examen_word(self, numero_tema):
        """
        Genera una versión de examen en formato Word
        
        Args:
            numero_tema (int): Número del tema a generar (1, 2, 3, ...)
            
        Returns:
            str: Ruta del archivo Word generado
        """
        # Generar letra del tema (A, B, C, ...)
        letra_tema = chr(64 + numero_tema)  # 65 es el código ASCII de 'A'
        nombre_archivo = os.path.join(self.directorio_examenes, f"Examen_Tema_{letra_tema}.docx")
        
        # Generar examen aleatorio
        preguntas_examen = self.pregunta_dao.generar_examen_aleatorio()
        
        # Crear Word
        if self.generar_word(preguntas_examen, nombre_archivo, f"Tema {letra_tema}"):
            return nombre_archivo
        return None
    
    def generar_examenes(self, cantidad_temas, formato="pdf"):
        """
        Genera múltiples versiones de exámenes en el formato especificado
        
        Args:
            cantidad_temas (int): Número de temas diferentes a generar
            formato (str): Formato de los exámenes ("pdf" o "word")
            
        Returns:
            list: Lista con las rutas de los archivos generados
        """
        rutas_archivos = []
        
        for i in range(cantidad_temas):
            if formato.lower() == "pdf":
                ruta_archivo = self.generar_examen_pdf(i+1)
            elif formato.lower() == "word":
                ruta_archivo = self.generar_examen_word(i+1)
            else:
                # Formato no soportado
                continue
                
            if ruta_archivo:
                rutas_archivos.append(ruta_archivo)
        
        return rutas_archivos
    
    def generar_pdf(self, preguntas, ruta_archivo, titulo_examen):
        """
        Genera un archivo PDF con las preguntas del examen
        
        Args:
            preguntas (list): Lista de objetos Pregunta para el examen
            ruta_archivo (str): Ruta donde se guardará el archivo PDF
            titulo_examen (str): Título del examen (Tema A, Tema B, etc.)
            
        Returns:
            bool: True si el PDF se generó correctamente, False en caso contrario
        """
        try:
            # Configurar el documento
            doc = SimpleDocTemplate(
                ruta_archivo,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Estilos
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(
                name='TituloPrincipal',
                fontName='Helvetica-Bold',
                fontSize=16,
                alignment=1,  # Centrado
                spaceBefore=12
            ))
            styles.add(ParagraphStyle(
                name='Pregunta',
                fontName='Helvetica-Bold',
                fontSize=11,
                spaceAfter=6
            ))
            styles.add(ParagraphStyle(
                name='Alternativa',
                fontName='Helvetica',
                fontSize=10,
                leftIndent=20,
                spaceAfter=3
            ))
            
            # Contenido del documento
            contenido = []
            
            # Agregar portada
            self.agregar_portada(contenido, styles, titulo_examen)
            
            # Agregar título
            contenido.append(Paragraph(f"EXAMEN DE ADMISIÓN - {titulo_examen}", styles['TituloPrincipal']))
            contenido.append(Spacer(1, 12))
            
            # Agregar instrucciones
            contenido.append(Paragraph(
                "Instrucciones: Marque la alternativa correcta para cada pregunta.",
                styles['Italic']
            ))
            contenido.append(Spacer(1, 12))
            
            # Agregar preguntas
            for i, pregunta in enumerate(preguntas):
                # Número y enunciado de la pregunta
                contenido.append(Paragraph(
                    f"{i+1}. {pregunta.enunciado}",
                    styles['Pregunta']
                ))
                
                # Alternativas
                contenido.append(Paragraph(f"a) {pregunta.alternativa_a}", styles['Alternativa']))
                contenido.append(Paragraph(f"b) {pregunta.alternativa_b}", styles['Alternativa']))
                contenido.append(Paragraph(f"c) {pregunta.alternativa_c}", styles['Alternativa']))
                contenido.append(Paragraph(f"d) {pregunta.alternativa_d}", styles['Alternativa']))
                contenido.append(Paragraph(f"e) {pregunta.alternativa_e}", styles['Alternativa']))
                
                contenido.append(Spacer(1, 12))
            
            # Construir el documento
            doc.build(contenido)
            return True
            
        except Exception as e:
            print(f"Error al generar el PDF: {e}")
            return False    
    def agregar_portada(self, contenido, styles, titulo_examen):
        """
        Agrega una portada al documento PDF
        """
        # Crear estilos específicos para la portada
        styles.add(ParagraphStyle(
            name='UniversidadTitulo',
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=30,
            alignment=1,
            spaceAfter=30
        ))
        
        styles.add(ParagraphStyle(
            name='ExamenTitulo',
            fontName='Helvetica-Bold',
            fontSize=26,
            alignment=1,
            leading=30,
        ))

        styles.add(ParagraphStyle(
            name='Tema',
            fontName='Helvetica-Bold',
            fontSize=32,
            alignment=1,
            leading=30,
            spaceBefore=20,
            spaceAfter=20
        ))

        # Nombre de la universidad (en dos líneas)
        contenido.append(Paragraph(
            'UNIVERSIDAD NACIONAL "SAN LUIS<br/>GONZAGA"',
            styles['UniversidadTitulo']
        ))
        
        # Examen de admisión
        contenido.append(Paragraph(
            "EXAMEN DE ADMISIÓN 2025",
            styles['ExamenTitulo']
        ))
        
        # Espacio antes del logo
        contenido.append(Spacer(1, 0))
        
        # Logo
        logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'images', 'logo.png')
        if os.path.exists(logo_path):
            img = Image(logo_path)
            img.drawHeight = 5*inch
            img.drawWidth = 5*inch
            contenido.append(img)
        
        # Espacio después del logo
        contenido.append(Spacer(1, 0))
        
        # Modalidad
        contenido.append(Paragraph(
            "MODALIDAD",
            styles['ExamenTitulo']
        ))
        contenido.append(Paragraph(
            "ORDINARIO",
            styles['ExamenTitulo'],
        ))
        
        # Espacio antes del tema
        contenido.append(Spacer(1, 0))
        
        # Tema
        contenido.append(Paragraph(
            f"TEMA: ({titulo_examen[-1]})",
            styles['Tema']
        ))
        
        # Espacio antes del pie de página
        contenido.append(Spacer(1, 0))
        
        # Pie de página
        contenido.append(Paragraph(
            "UNIVERSIDAD LICENCIADA POR SUNEDU",
            ParagraphStyle(
                'PiePagina',
                parent=styles['Normal'],
                alignment=1,
                spaceBefore=14,
                fontSize=16
            )
        ))
        
        # Salto de página
        contenido.append(Paragraph("<br clear=all style='page-break-before:always'/>", styles['Normal']))
        
    def generar_word(self, preguntas, ruta_archivo, titulo_examen):
        """
        Genera un archivo Word con las preguntas del examen
        
        Args:
            preguntas (list): Lista de objetos Pregunta para el examen
            ruta_archivo (str): Ruta donde se guardará el archivo Word
            titulo_examen (str): Título del examen (Tema A, Tema B, etc.)
            
        Returns:
            bool: True si el Word se generó correctamente, False en caso contrario
        """
        try:
            # Crear un nuevo documento
            doc = Document()
            
            # Configurar estilos
            styles = doc.styles
            
            # Estilo para título principal
            titulo_style = styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
            titulo_style.font.name = 'Arial'
            titulo_style.font.size = Pt(16)
            titulo_style.font.bold = True
            
            # Estilo para preguntas
            pregunta_style = styles.add_style('Pregunta', WD_STYLE_TYPE.PARAGRAPH)
            pregunta_style.font.name = 'Arial'
            pregunta_style.font.size = Pt(11)
            pregunta_style.font.bold = True
            
            # Estilo para alternativas
            alternativa_style = styles.add_style('Alternativa', WD_STYLE_TYPE.PARAGRAPH)
            alternativa_style.font.name = 'Arial'
            alternativa_style.font.size = Pt(10)
            alternativa_style.paragraph_format.left_indent = Inches(0.25)
            
            # Agregar portada
            self._agregar_portada_word(doc, titulo_examen)
            
            # Agregar título
            titulo = doc.add_paragraph(f"EXAMEN DE ADMISIÓN - {titulo_examen}", 'TituloPrincipal')
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
            # Agregar instrucciones
            instrucciones = doc.add_paragraph("Instrucciones: Marque la alternativa correcta para cada pregunta.")
            # Crear o configurar estilo Italic
            if 'Italic' not in styles:
                italic_style = styles.add_style('Italic', WD_STYLE_TYPE.PARAGRAPH)
                italic_style.font.name = 'Arial'
                italic_style.font.size = Pt(10)
                italic_style.font.italic = True
            else:
                # Asegurar que el estilo tenga la propiedad italic configurada
                styles['Italic'].font.italic = True
                styles['Italic'].font.name = 'Arial'
                styles['Italic'].font.size = Pt(10)
            instrucciones.style = styles['Italic']
            doc.add_paragraph()
            
            # Agregar preguntas
            for i, pregunta in enumerate(preguntas):
                # Número y enunciado de la pregunta
                p = doc.add_paragraph(f"{i+1}. {pregunta.enunciado}", 'Pregunta')
                
                # Alternativas
                doc.add_paragraph(f"a) {pregunta.alternativa_a}", 'Alternativa')
                doc.add_paragraph(f"b) {pregunta.alternativa_b}", 'Alternativa')
                doc.add_paragraph(f"c) {pregunta.alternativa_c}", 'Alternativa')
                doc.add_paragraph(f"d) {pregunta.alternativa_d}", 'Alternativa')
                doc.add_paragraph(f"e) {pregunta.alternativa_e}", 'Alternativa')
                
                doc.add_paragraph()
            
            # Guardar el documento
            doc.save(ruta_archivo)
            return True
            
        except Exception as e:
            print(f"Error al generar el Word: {e}")
            return False
    
    def _agregar_portada_word(self, doc, titulo_examen):
        """
        Agrega una portada al documento Word
        
        Args:
            doc (Document): Documento Word
            titulo_examen (str): Título del examen (Tema A, Tema B, etc.)
        """
        # Crear estilos para la portada
        styles = doc.styles
        
        # Estilo para título de universidad
        univ_style = styles.add_style('UniversidadTitulo', WD_STYLE_TYPE.PARAGRAPH)
        univ_style.font.name = 'Arial'
        univ_style.font.size = Pt(16)
        univ_style.font.bold = True
        
        # Estilo para título de examen
        examen_style = styles.add_style('ExamenTitulo', WD_STYLE_TYPE.PARAGRAPH)
        examen_style.font.name = 'Arial'
        examen_style.font.size = Pt(24)
        examen_style.font.bold = True
        
        # Estilo para tema
        tema_style = styles.add_style('Tema', WD_STYLE_TYPE.PARAGRAPH)
        tema_style.font.name = 'Arial'
        tema_style.font.size = Pt(32)
        tema_style.font.bold = True
        
        # Nombre de la universidad
        univ = doc.add_paragraph('UNIVERSIDAD NACIONAL "SAN LUIS GONZAGA"', 'UniversidadTitulo')
        univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Examen de admisión
        examen = doc.add_paragraph("EXAMEN DE ADMISIÓN 2025", 'ExamenTitulo')
        examen.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Logo
        logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'images', 'logo.png')
        if os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        # Modalidad
        modalidad_title = doc.add_paragraph("MODALIDAD", 'ExamenTitulo')
        modalidad_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        modalidad = doc.add_paragraph("ORDINARIO", 'ExamenTitulo')
        modalidad.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tema
        tema = doc.add_paragraph(f"TEMA: ({titulo_examen[-1]})", 'Tema')
        tema.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espacio antes del pie de página
        doc.add_paragraph()
        
        # Pie de página
        # Crear estilo para pie de página si no existe
        if 'PiePagina' not in styles:
            pie_style = styles.add_style('PiePagina', WD_STYLE_TYPE.PARAGRAPH)
            pie_style.font.name = 'Arial'
            pie_style.font.size = Pt(16)
            pie_style.font.bold = True
        
        pie = doc.add_paragraph("UNIVERSIDAD LICENCIADA POR SUNEDU")
        pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pie.style = styles['PiePagina']